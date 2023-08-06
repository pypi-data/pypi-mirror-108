import csv
import cv2
import docx
from docx import Document
from docxcompose.composer import Composer
import numpy as np
import os
import random
from random import randint
import string



def random_drawing():
    """
    This function generates a random picture from geometric shapes
    """
    document = Document()
    document.save('test.docx')

    # Create a black image
    img = np.zeros((512, 512, 3), np.uint8)

    # Draw a diagonal blue line with thickness of 5 px
    cv2.line(img, (0, 0), (randint(1, 511), randint(1, 511)), (randint(1, 511), 0, 0), 5)

    cv2.rectangle(img, (randint(1, 511), randint(1, 511)), (randint(1, 511), randint(1, 511)),
                      (randint(1, 511), randint(1, 511), randint(1, 511)), 3)
    cv2.circle(img, (randint(1, 511), randint(1, 511)), 63, (0, 0, randint(1, 511)), -1)
    cv2.ellipse(img, (randint(1, 511), randint(1, 511)), (randint(1, 511), randint(1, 511)), 0, 0, 180, 255, -1)

    pts = np.array([[10, 5], [20, 30], [70, 20], [50, 10]], np.int32)
    pts = pts.reshape((-1, 1, 2))
    cv2.polylines(img, [pts], True, (randint(1, 511), randint(1, 511), randint(1, 511)))

    cv2.waitKey(0)
    cv2.imwrite('Image.jpg', img)
    doc = docx.Document()

    doc.add_paragraph('Random generated picture')
    doc.add_picture('Image.jpg', width=docx.shared.Cm(15))
    document.add_page_break()
    doc.save('test.docx')

def get_table():
    """
    This function generates a random fixed-sized table of numbers
    :return:
    """
    file = open("test1.txt", "w")
    file.close()

    def get_string():
        random_size = random.randint(5, 5)
        return [[0 for _ in range(random_size)] for _ in range(random_size)]

    my_array = np.random.rand(10, 4)
    np.savetxt('test1.txt', my_array, fmt='%4.6f', delimiter=' ')
    with open('test1.txt') as infile, open('test1.csv', 'w') as outfile:
        for line in infile:
            outfile.write(line.replace(' ', ','))
    doc = docx.Document()

    with open('test1.csv', newline='') as f:
        csv_reader = csv.reader(f)
        doc.add_page_break()
        doc.add_paragraph('Random generated table')

        csv_headers = next(csv_reader)
        csv_cols = len(csv_headers)

        table = doc.add_table(rows=2, cols=csv_cols)
        hdr_cells = table.rows[0].cells

        for i in range(csv_cols):
            hdr_cells[i].text = csv_headers[i]

        for row in csv_reader:
            row_cells = table.add_row().cells
            for i in range(csv_cols):
                row_cells[i].text = row[i]

    doc.save("test1.docx")
def merge(path:str):
    """
    This function combines several word files into one single file at the specified path. Possibly can be used as a standalone helper library or smth idk ¯\_(ツ)_/¯
    :param path: selected save path
    :return:
    """
    files=["test.docx", "test1.docx", "test2.docx"]
    new_document = Document()
    composer = Composer(new_document)
    for fn in files:
        composer.append(Document(fn))
    composer.save(path)
    print('Merged file just saved to ', path, '!')

def trash_clean():
    """
    This function cleans up intermediate files. In fact, they can remove this program too
    """
    path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'test.txt')
    os.remove(path)
    path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'test1.csv')
    os.remove(path)
    path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'test1.txt')
    os.remove(path)
    path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'test.docx')
    os.remove(path)
    path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'test1.docx')
    os.remove(path)
    path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'test2.docx')
    os.remove(path)
    path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'Image.jpg')
    os.remove(path)
def random_string(length: int):
    file = open("test.txt", "w")
    file.close()
    if length < 0:
        length=0

def generate_random_string(length):
    """
    This function generates a random set of letters of a given length
    :param length: length- number of letters of random text
    :return:
    """
    letters = string.ascii_lowercase
    rand_string = ''.join(random.choice(letters) for i in range(length))
    file = open('test.txt', 'w')
    print("Random string of length", length, "is:", rand_string, file=file)
    file.close()

    generate_random_string(length)

    file = open("test.txt", "r")
    data = file.read()
    doc = docx.Document()
    doc.add_page_break()
    doc.add_paragraph('Random generated string:')
    doc.add_paragraph(data)
    doc.save('test2.docx')
    file.close()
    return length
